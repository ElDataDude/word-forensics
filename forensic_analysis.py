#!/usr/bin/env python3

import argparse
import json
import os
import sys
import zipfile
from datetime import datetime
from pathlib import Path
import xml.etree.ElementTree as ET
import binascii
import difflib
from openai import OpenAI
import logging
import time
import hashlib
from docx import Document
from typing import Dict, List, Tuple, Any
from dotenv import load_dotenv
from statistical_analysis import ForensicStatisticalAnalyzer

# Set up logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

class WordForensicAnalyzer:
    def __init__(self, target_path: str, same_origin_path: str, reference_path: str):
        """Initialize the analyzer with paths to documents for comparison."""
        load_dotenv()
        
        # Initialize OpenAI client with minimal configuration
        api_key = os.getenv('OPENAI_API_KEY')
        if not api_key:
            raise ValueError("OPENAI_API_KEY environment variable is not set")
            
        self.client = OpenAI(
            api_key=api_key,
            base_url="https://api.openai.com/v1",  # Explicitly set the base URL
            timeout=60.0,  # Set a reasonable timeout
            max_retries=3  # Set max retries
        )
        
        self.target_path = Path(target_path)
        self.same_origin_path = Path(same_origin_path)
        self.reference_path = Path(reference_path)
        
        # Create output directory within project structure
        self.output_dir = Path(os.path.dirname(os.path.abspath(__file__))) / "output"
        self.output_dir.mkdir(parents=True, exist_ok=True)
        
        # Cache directory for storing intermediate results
        self.cache_dir = self.output_dir / ".cache"
        self.cache_dir.mkdir(parents=True, exist_ok=True)
        
        # Initialize statistical analyzer
        self.statistical_analyzer = ForensicStatisticalAnalyzer(
            reference_dir=self.reference_path.parent,
            cache_dir=self.cache_dir
        )
        self.statistical_analyzer.set_analyzer(self)
        
    def validate_files(self) -> None:
        """Validate existence and format of input files."""
        # Validate input files
        for file_path in [self.target_path, self.same_origin_path, self.reference_path]:
            if not file_path.exists():
                raise FileNotFoundError(f"File not found: {file_path}")
            if file_path.suffix.lower() != '.docx':
                raise ValueError(f"Invalid file format for {file_path}. Must be .docx")
        
    def extract_metadata(self, docx_path: Path) -> Dict[str, Any]:
        """Extract metadata from a Word document."""
        try:
            doc = Document(docx_path)
            if not hasattr(doc, 'core_properties'):
                logging.warning(f"No core properties found in {docx_path.name}")
                return {}
                
            core_props = doc.core_properties
            
            # Helper function to safely get attributes
            def get_prop(obj, attr):
                try:
                    val = getattr(obj, attr, None)
                    # Convert datetime objects to string
                    if hasattr(val, 'isoformat'):
                        return val.isoformat()
                    return val
                except Exception as e:
                    logging.warning(f"Error getting property {attr}: {str(e)}")
                    return None
            
            metadata = {
                "author": get_prop(core_props, 'author'),
                "last_modified_by": get_prop(core_props, 'last_modified_by'),
                "created": get_prop(core_props, 'created'),
                "modified": get_prop(core_props, 'modified'),
                "revision": get_prop(core_props, 'revision'),
                "title": get_prop(core_props, 'title'),
                "category": get_prop(core_props, 'category'),
                "keywords": get_prop(core_props, 'keywords'),
                "comments": get_prop(core_props, 'comments'),
                "identifier": get_prop(core_props, 'identifier'),
                "language": get_prop(core_props, 'language'),
                "subject": get_prop(core_props, 'subject'),
                "version": get_prop(core_props, 'version')
            }
            
            # Remove None values for cleaner output
            return {k: v for k, v in metadata.items() if v is not None}
            
        except Exception as e:
            logging.error(f"Error extracting metadata from {docx_path.name}: {str(e)}")
            return {}

    def analyze_content(self, docx_path: Path) -> Dict[str, Any]:
        """Analyze document content including text, styles, and structure."""
        try:
            doc = Document(docx_path)
            
            content_analysis = {
                "paragraphs": 0,
                "sections": 0,
                "styles": [],
                "text_content": "",
                "fonts_used": set(),
                "tracked_changes": [],
                "comments": [],
                "embedded_objects": []
            }
            
            # Safely get document structure
            try:
                content_analysis["paragraphs"] = len(doc.paragraphs)
            except:
                logging.warning(f"Could not count paragraphs in {docx_path.name}")
                
            try:
                content_analysis["sections"] = len(doc.sections)
            except:
                logging.warning(f"Could not count sections in {docx_path.name}")
            
            # Analyze text and styles
            for paragraph in doc.paragraphs:
                try:
                    content_analysis["text_content"] += paragraph.text + "\n"
                    if hasattr(paragraph, 'style') and paragraph.style and paragraph.style.name:
                        if paragraph.style.name not in content_analysis["styles"]:
                            content_analysis["styles"].append(paragraph.style.name)
                    
                    # Extract font information from runs
                    for run in paragraph.runs:
                        if hasattr(run, 'font') and hasattr(run.font, 'name') and run.font.name:
                            content_analysis["fonts_used"].add(run.font.name)
                except Exception as e:
                    logging.warning(f"Error processing paragraph in {docx_path.name}: {str(e)}")
                    continue
            
            # Convert set to list for JSON serialization
            content_analysis["fonts_used"] = list(content_analysis["fonts_used"])
            return content_analysis
            
        except Exception as e:
            logging.error(f"Error analyzing content of {docx_path.name}: {str(e)}")
            return {
                "paragraphs": 0,
                "sections": 0,
                "styles": [],
                "text_content": "",
                "fonts_used": [],
                "tracked_changes": [],
                "comments": [],
                "embedded_objects": [],
                "error": str(e)
            }

    def analyze_ooxml_structure(self, docx_path: Path) -> Dict[str, Any]:
        """Analyze the internal OOXML structure of the document."""
        structure_info = {
            "xml_files": [],
            "relationships": [],
            "content_types": [],
            "element_counts": {},
            "custom_properties": [],
            "error": None
        }
        
        try:
            with zipfile.ZipFile(docx_path, 'r') as zip_ref:
                # Analyze XML files
                for item in zip_ref.namelist():
                    if item.endswith('.xml') or item.endswith('.rels'):
                        try:
                            with zip_ref.open(item) as xml_file:
                                tree = ET.parse(xml_file)
                                root = tree.getroot()
                                
                                structure_info["xml_files"].append({
                                    "path": item,
                                    "root_tag": root.tag,
                                    "namespace": root.tag.split('}')[0].strip('{') if '}' in root.tag else None,
                                    "child_count": len(list(root))
                                })
                                
                                # Count elements
                                element_counts = {}
                                for elem in root.iter():
                                    tag = elem.tag.split('}')[-1] if '}' in elem.tag else elem.tag
                                    element_counts[tag] = element_counts.get(tag, 0) + 1
                                structure_info["element_counts"][item] = element_counts
                        except Exception as e:
                            logging.warning(f"Error parsing XML file {item} in {docx_path.name}: {str(e)}")
                            continue
                
                # Analyze relationships
                if '[Content_Types].xml' in zip_ref.namelist():
                    try:
                        with zip_ref.open('[Content_Types].xml') as types_xml:
                            types_tree = ET.parse(types_xml)
                            types_root = types_tree.getroot()
                            for override in types_root.findall('{http://schemas.openxmlformats.org/package/2006/content-types}Override'):
                                structure_info["content_types"].append({
                                    "part_name": override.get('PartName'),
                                    "content_type": override.get('ContentType')
                                })
                    except Exception as e:
                        logging.warning(f"Error parsing content types in {docx_path.name}: {str(e)}")
                
                # Analyze custom properties
                if 'docProps/custom.xml' in zip_ref.namelist():
                    try:
                        with zip_ref.open('docProps/custom.xml') as custom_xml:
                            custom_tree = ET.parse(custom_xml)
                            custom_root = custom_tree.getroot()
                            for prop in custom_root.findall('.//{http://schemas.openxmlformats.org/officeDocument/2006/custom-properties}property'):
                                structure_info["custom_properties"].append({
                                    "name": prop.get('name'),
                                    "type": next(iter(prop)).tag.split('}')[-1] if len(prop) > 0 else None,
                                    "value": next(iter(prop)).text if len(prop) > 0 else None
                                })
                    except Exception as e:
                        logging.warning(f"Error parsing custom properties in {docx_path.name}: {str(e)}")
            
            return structure_info
            
        except Exception as e:
            error_msg = f"Error analyzing OOXML structure of {docx_path.name}: {str(e)}"
            logging.error(error_msg)
            structure_info["error"] = error_msg
            return structure_info

    def analyze_binary_content(self, docx_path: Path) -> Dict[str, Any]:
        """Perform binary analysis of the document."""
        binary_info = {
            "file_size": 0,
            "readable_strings": [],
            "user_paths": [],
            "template_paths": [],
            "system_markers": [],
            "binary_signatures": [],
            "error": None
        }
        
        try:
            # Get file size
            binary_info["file_size"] = os.path.getsize(docx_path)
            
            # Read file in binary mode
            with open(docx_path, 'rb') as f:
                content = f.read()
                
                # Convert to hex for signature analysis
                hex_content = binascii.hexlify(content[:4096]).decode('utf-8')  # Analyze first 4KB for signatures
                
                # Look for common binary signatures
                signatures = {
                    '504B0304': 'ZIP Archive (DOCX)',
                    '504B0506': 'ZIP End of Central Directory',
                    '504B0708': 'ZIP Data Descriptor'
                }
                
                for sig, desc in signatures.items():
                    if sig.lower() in hex_content.lower():
                        binary_info["binary_signatures"].append({
                            "signature": sig,
                            "description": desc,
                            "offset": hex_content.lower().index(sig.lower()) // 2
                        })
                
                # Extract readable strings (at least 4 chars long)
                current_string = ""
                for byte in content:
                    char = chr(byte)
                    if 32 <= byte <= 126:  # Printable ASCII
                        current_string += char
                    else:
                        if len(current_string) >= 4:
                            # Filter out common noise
                            if not any(noise in current_string.lower() for noise in ['xmlns', 'http://', 'https://']):
                                s = current_string
                                s_lower = s.lower()
                                
                                # User directory paths
                                if '/users/' in s_lower or r'c:\users' in s_lower:
                                    binary_info["user_paths"].append(s)
                                
                                # Template paths
                                elif 'template' in s_lower and ('.dot' in s_lower or '.dotx' in s_lower):
                                    binary_info["template_paths"].append(s)
                                
                                # System markers
                                elif any(marker in s_lower for marker in ['windows', 'microsoft', 'office', 'word']):
                                    binary_info["system_markers"].append(s)
                                
                                # Other readable strings
                                else:
                                    binary_info["readable_strings"].append(s)
                        
                        current_string = ""
            
            # Limit the number of strings to avoid excessive output
            max_strings = 100
            binary_info["readable_strings"] = binary_info["readable_strings"][:max_strings]
            
            return binary_info
            
        except Exception as e:
            error_msg = f"Error analyzing binary content of {docx_path.name}: {str(e)}"
            logging.error(error_msg)
            binary_info["error"] = error_msg
            return binary_info

    def find_origin_evidence(self) -> Dict[str, Any]:
        """Find definitive evidence of shared origin between target and same-origin files."""
        target_binary = self.analyze_binary_content(self.target_path)
        same_origin_binary = self.analyze_binary_content(self.same_origin_path)
        reference_binary = self.analyze_binary_content(self.reference_path)
        
        target_meta = self.extract_metadata(self.target_path)
        same_origin_meta = self.extract_metadata(self.same_origin_path)
        reference_meta = self.extract_metadata(self.reference_path)
        
        target_content = self.analyze_content(self.target_path)
        same_origin_content = self.analyze_content(self.same_origin_path)
        reference_content = self.analyze_content(self.reference_path)
        
        evidence = {
            "definitive_markers": [],
            "strong_indicators": [],
            "potential_indicators": [],
            "metadata_matches": {},
            "binary_signatures": [],
            "shared_paths": {
                "user_paths": [],
                "template_paths": [],
                "system_markers": []
            }
        }
        
        # Helper function to find matching paths
        def find_matching_paths(list1, list2, ref_list):
            shared = set()
            for path1 in list1:
                for path2 in list2:
                    # If paths match and aren't in reference file
                    if path1.lower() == path2.lower() and not any(path1.lower() == ref.lower() for ref in ref_list):
                        shared.add(path1)
            return list(shared)
        
        # Helper function to check metadata matches
        def check_metadata_match(field: str, target_val, same_origin_val, ref_val) -> bool:
            if not target_val or not same_origin_val:
                return False
            return (
                str(target_val).lower() == str(same_origin_val).lower() and
                (not ref_val or str(ref_val).lower() != str(target_val).lower())
            )
        
        # Check metadata matches
        metadata_fields = {
            "author": "Author name",
            "last_modified_by": "Last modified by",
            "title": "Document title",
            "template": "Template name",
            "company": "Company name",
            "application": "Application name",
            "app_version": "Application version"
        }
        
        for field, description in metadata_fields.items():
            if check_metadata_match(
                field,
                target_meta.get(field),
                same_origin_meta.get(field),
                reference_meta.get(field)
            ):
                evidence["metadata_matches"][field] = {
                    "value": target_meta[field],
                    "description": description
                }
                
                # Author and last_modified_by are strong indicators
                if field in ["author", "last_modified_by"]:
                    evidence["strong_indicators"].append({
                        "type": "metadata_match",
                        "field": field,
                        "value": target_meta[field],
                        "explanation": f"Identical {description} found in both target and same-origin files, but not in reference"
                    })
        
        # Find shared paths that aren't in reference file
        evidence["shared_paths"]["user_paths"] = find_matching_paths(
            target_binary["user_paths"],
            same_origin_binary["user_paths"],
            reference_binary["user_paths"]
        )
        
        evidence["shared_paths"]["template_paths"] = find_matching_paths(
            target_binary["template_paths"],
            same_origin_binary["template_paths"],
            reference_binary["template_paths"]
        )
        
        evidence["shared_paths"]["system_markers"] = find_matching_paths(
            target_binary["system_markers"],
            same_origin_binary["system_markers"],
            reference_binary["system_markers"]
        )
        
        # Analyze shared binary signatures
        target_sigs = set(sig["signature"] for sig in target_binary.get("binary_signatures", []))
        same_origin_sigs = set(sig["signature"] for sig in same_origin_binary.get("binary_signatures", []))
        ref_sigs = set(sig["signature"] for sig in reference_binary.get("binary_signatures", []))
        
        shared_sigs = target_sigs & same_origin_sigs - ref_sigs
        for sig in shared_sigs:
            evidence["binary_signatures"].append({
                "signature": sig,
                "explanation": "Identical binary signature found in both target and same-origin files, but not in reference"
            })
        
        # Analyze the evidence
        for path_type, paths in evidence["shared_paths"].items():
            for path in paths:
                path_lower = path.lower()
                
                # Definitive markers (user-specific paths with usernames)
                if ('/users/' in path_lower or r'c:\users' in path_lower) and any(
                    part for part in path_lower.split('/') if len(part) > 2
                    and not part.startswith('.')
                    and not any(common in part for common in ['public', 'shared', 'default', 'common'])
                ):
                    evidence["definitive_markers"].append({
                        "type": "user_specific_path",
                        "value": path,
                        "explanation": "Identical user-specific path with username found in both target and same-origin files, but not in reference"
                    })
                
                # Strong indicators (template paths, application paths)
                elif any(marker in path_lower for marker in [
                    'template', '.dotx', '.dotm', 'normal.dot',
                    'application data', 'appdata', 'microsoft/word'
                ]):
                    evidence["strong_indicators"].append({
                        "type": "shared_template",
                        "value": path,
                        "explanation": "Identical template or application-specific path found"
                    })
                
                # Potential indicators (system paths, temp paths)
                else:
                    evidence["potential_indicators"].append({
                        "type": "shared_path",
                        "value": path,
                        "explanation": "Identical system path or marker found"
                    })
        
        # Check for identical styles and fonts (potential indicators)
        shared_styles = set(target_content["styles"]) & set(same_origin_content["styles"]) - set(reference_content["styles"])
        if shared_styles:
            evidence["potential_indicators"].append({
                "type": "shared_styles",
                "value": list(shared_styles),
                "explanation": "Identical document styles found in both target and same-origin files, but not in reference"
            })
        
        shared_fonts = set(target_content["fonts_used"]) & set(same_origin_content["fonts_used"]) - set(reference_content["fonts_used"])
        if shared_fonts:
            evidence["potential_indicators"].append({
                "type": "shared_fonts",
                "value": list(shared_fonts),
                "explanation": "Identical fonts found in both target and same-origin files, but not in reference"
            })
        
        return evidence

    def compare_files(self) -> Dict[str, Any]:
        """Compare the target file with same-origin and reference files."""
        target_content = self.analyze_content(self.target_path)
        same_origin_content = self.analyze_content(self.same_origin_path)
        reference_content = self.analyze_content(self.reference_path)
        
        # Calculate similarity scores
        same_origin_similarity = difflib.SequenceMatcher(
            None, 
            target_content["text_content"], 
            same_origin_content["text_content"]
        ).ratio() * 100
        
        reference_similarity = difflib.SequenceMatcher(
            None, 
            target_content["text_content"], 
            reference_content["text_content"]
        ).ratio() * 100
        
        return {
            "similarity_scores": {
                "same_origin": round(same_origin_similarity, 2),
                "reference": round(reference_similarity, 2)
            },
            "shared_styles": {
                "same_origin": list(set(target_content["styles"]) & set(same_origin_content["styles"])),
                "reference": list(set(target_content["styles"]) & set(reference_content["styles"]))
            },
            "shared_fonts": {
                "same_origin": list(set(target_content["fonts_used"]) & set(same_origin_content["fonts_used"])),
                "reference": list(set(target_content["fonts_used"]) & set(reference_content["fonts_used"]))
            }
        }

    def generate_report(self) -> Tuple[Dict[str, Any], str]:
        """Generate the complete analysis report and summary."""
        report = {
            "analysis_timestamp": datetime.now().isoformat(),
            "target_file": str(self.target_path),
            "same_origin_file": str(self.same_origin_path),
            "reference_file": str(self.reference_path),
            "metadata_analysis": {
                "target": self.extract_metadata(self.target_path),
                "same_origin": self.extract_metadata(self.same_origin_path),
                "reference": self.extract_metadata(self.reference_path)
            },
            "content_analysis": {
                "target": self.analyze_content(self.target_path),
                "comparisons": self.compare_files()
            },
            "ooxml_analysis": {
                "target": self.analyze_ooxml_structure(self.target_path),
                "same_origin": self.analyze_ooxml_structure(self.same_origin_path),
                "reference": self.analyze_ooxml_structure(self.reference_path)
            },
            "binary_analysis": {
                "target": self.analyze_binary_content(self.target_path),
                "same_origin": self.analyze_binary_content(self.same_origin_path),
                "reference": self.analyze_binary_content(self.reference_path)
            },
            "origin_evidence": self.find_origin_evidence()
        }
        
        # Add statistical analysis
        try:
            target_data = {
                "metadata": report["metadata_analysis"]["target"],
                "content": report["content_analysis"]["target"],
                "ooxml": report["ooxml_analysis"]["target"],
                "binary": report["binary_analysis"]["target"]
            }
            same_origin_data = {
                "metadata": report["metadata_analysis"]["same_origin"],
                "content": self.analyze_content(self.same_origin_path),  # Get content directly
                "ooxml": report["ooxml_analysis"]["same_origin"],
                "binary": report["binary_analysis"]["same_origin"]
            }
            statistical_results = self.statistical_analyzer.analyze_similarity(
                target_data=target_data,
                same_origin_data=same_origin_data
            )
            report["statistical_analysis"] = statistical_results
        except Exception as e:
            logging.error(f"Error performing statistical analysis: {str(e)}")
            report["statistical_analysis"] = {"error": str(e)}
        
        # Generate natural language summary
        summary = self._generate_summary(report)
        
        return report, summary

    def _get_cache_key(self, report: Dict[str, Any]) -> str:
        """Generate a cache key based on the report content."""
        # Only use relevant parts of the report for the cache key
        key_data = {
            'evidence': report['origin_evidence'],
            'metadata': report['metadata_analysis'],
            'binary': report['binary_analysis']
        }
        return hashlib.md5(json.dumps(key_data, sort_keys=True).encode()).hexdigest()

    def _get_cached_summary(self, cache_key: str) -> str:
        """Try to get a cached summary."""
        cache_file = self.cache_dir / f"{cache_key}.txt"
        if cache_file.exists():
            logging.info("Using cached summary")
            return cache_file.read_text()
        return None

    def _cache_summary(self, cache_key: str, summary: str) -> None:
        """Cache a successful summary."""
        cache_file = self.cache_dir / f"{cache_key}.txt"
        cache_file.write_text(summary)
        logging.info("Cached summary for future use")

    def _generate_summary(self, report: Dict[str, Any]) -> str:
        """Generate a natural language summary using OpenAI API with caching and improved retry logic."""
        # Try to get cached summary first
        cache_key = self._get_cache_key(report)
        cached_summary = self._get_cached_summary(cache_key)
        if cached_summary:
            return cached_summary

        evidence = report["origin_evidence"]
        statistical = report.get("statistical_analysis", {})
        
        # Create a more concise prompt
        prompt = (
            "Analyze Word document forensics evidence:\n"
            f"Definitive markers: {len(evidence.get('definitive_markers', []))}\n"
            f"Strong indicators: {len(evidence.get('strong_indicators', []))}\n"
            f"Potential indicators: {len(evidence.get('potential_indicators', []))}\n\n"
            "Statistical Analysis:\n"
            f"{json.dumps(statistical.get('statistical_summary', {}), indent=2)}\n\n"
            "Key metadata matches:\n"
            f"{json.dumps(evidence.get('metadata_matches', {}), indent=2)}\n\n"
            "Binary signatures:\n"
            f"{json.dumps(evidence.get('binary_signatures', {}), indent=2)}\n\n"
            "Provide:\n"
            "1. Conclusion on shared origin\n"
            "2. Evidence strength (definitive/strong/potential)\n"
            "3. Key technical markers\n"
            "4. Statistical confidence\n"
            "5. Caveats\n"
            "6. Further investigation needs"
        )

        max_retries = 5
        base_delay = 60  # 1 minute
        
        for attempt in range(max_retries):
            try:
                print(f"\nAttempting to generate AI summary (attempt {attempt + 1}/{max_retries})...")
                response = self.client.chat.completions.create(
                    model="gpt-4",
                    messages=[
                        {"role": "system", "content": "You are a digital forensics expert analyzing Word documents for shared origin evidence."},
                        {"role": "user", "content": prompt}
                    ],
                    max_tokens=800,
                    temperature=0.3
                )
                summary = response.choices[0].message.content
                # Cache successful response
                self._cache_summary(cache_key, summary)
                logging.info("AI summary generated successfully")
                return summary
            except Exception as e:
                if "rate limit" in str(e).lower():
                    if attempt < max_retries - 1:
                        wait_time = base_delay * (2 ** attempt)  # Exponential backoff
                        logging.warning(f"Rate limit hit. Waiting {wait_time} seconds before retry {attempt + 1}/{max_retries}")
                        print(f"\nRate limit reached. Waiting {wait_time} seconds before retrying...")
                        print(f"(You can press Ctrl+C to skip to template-based summary)")
                        try:
                            time.sleep(wait_time)
                            continue
                        except KeyboardInterrupt:
                            print("\nSkipping to template-based summary...")
                            break
                logging.error(f"Error generating AI summary: {str(e)}")
                print(f"\nError generating AI summary: {str(e)}")
                break

        print("Falling back to template-based summary...")
        return self._generate_template_summary(report)

    def _generate_template_summary(self, report: Dict[str, Any]) -> str:
        """Fallback template-based summary generator."""
        evidence = report["origin_evidence"]
        statistical = report.get("statistical_analysis", {})
        
        summary_lines = []
        
        # Add statistical analysis results if available
        if "statistical_summary" in statistical:
            stats = statistical["statistical_summary"]
            summary_lines.extend([
                "Statistical Analysis Results:",
                f"- Similarity Percentile: {stats['similarity_percentile']:.1f}%",
                f"- Z-Score: {stats['z_score']:.2f}",
                f"- Likelihood Ratio: {stats['likelihood_ratio']:.2f}",
                f"- Reference Sample Size: {stats['reference_sample_size']}",
                "",
                "Statistical Interpretation:",
                statistical["interpretation"]["percentile_interpretation"],
                statistical["interpretation"]["z_score_interpretation"],
                statistical["interpretation"]["likelihood_interpretation"],
                statistical["interpretation"]["confidence_note"] if statistical["interpretation"]["confidence_note"] else "",
                ""
            ])

        # Add evidence summary
        summary_lines.extend([
            "1. Conclusion on shared origin:",
            "Based on the available evidence, " + (
                "there is definitive proof" if evidence.get("definitive_markers", []) else
                "there are strong indicators" if evidence.get("strong_indicators", []) else
                "it is possible"
            ) + " that the Word documents share a common origin. " +
            f"Found {len(evidence.get('definitive_markers', []))} definitive markers, " +
            f"{len(evidence.get('strong_indicators', []))} strong indicators, and " +
            f"{len(evidence.get('potential_indicators', []))} potential indicators.",
            "",
            "2. Evidence strength:",
            "The evidence strength is " + (
                "definitive" if evidence.get("definitive_markers", []) else
                "strong" if evidence.get("strong_indicators", []) else
                "potential to strong" if len(evidence.get("potential_indicators", [])) > 10 else
                "potential"
            ) + ". " + (
                "There are definitive markers that prove shared origin." if evidence.get("definitive_markers", []) else
                "There are strong indicators but no definitive proof." if evidence.get("strong_indicators", []) else
                f"There are {len(evidence.get('potential_indicators', []))} potential indicators suggesting shared origin."
            ),
            "",
            "3. Key technical markers:",
            "The key technical markers in this case are:",
        ])
        
        # Add definitive markers
        for marker in evidence.get("definitive_markers", []):
            summary_lines.append(f"- [DEFINITIVE] {marker['explanation']}")
        
        # Add strong indicators
        for indicator in evidence.get("strong_indicators", []):
            summary_lines.append(f"- [STRONG] {indicator['explanation']}")
        
        # Add top potential indicators (limit to 5)
        top_potential = evidence.get("potential_indicators", [])[:5]
        if top_potential:
            summary_lines.append("\nTop potential indicators:")
            for indicator in top_potential:
                summary_lines.append(f"- {indicator['explanation']}")
        
        # Add statistical confidence
        if "statistical_summary" in statistical:
            stats = statistical["statistical_summary"]
            confidence = (
                "very high" if stats["likelihood_ratio"] > 0.9 else
                "high" if stats["likelihood_ratio"] > 0.7 else
                "moderate" if stats["likelihood_ratio"] > 0.5 else
                "low"
            )
            summary_lines.extend([
                "",
                "4. Statistical confidence:",
                f"Based on statistical analysis, there is a {confidence} level of confidence " +
                f"({stats['likelihood_ratio']*100:.1f}%) that the documents share a common origin. " +
                f"This is supported by a similarity score in the {stats['similarity_percentile']:.1f}th percentile " +
                f"compared to the reference set of {stats['reference_sample_size']} documents."
            ])
        else:
            summary_lines.extend([
                "",
                "4. Statistical confidence:",
                "Statistical analysis was not available for this comparison."
            ])
        
        # Add caveats and recommendations
        summary_lines.extend([
            "",
            "5. Caveats:",
            "- Metadata can be manipulated and should not be solely relied upon",
            "- Statistical analysis is based on available reference documents",
            "- Binary signatures and paths may be system-dependent",
            "",
            "6. Further investigation needs:",
            "- Expand reference document dataset for more robust statistical analysis",
            "- Perform deeper analysis of document structure and formatting",
            "- Consider timeline analysis of document modifications",
            "- Analyze any embedded objects or macros"
        ])
        
        return "\n".join(summary_lines)

def find_docx_files(directory: Path) -> List[Path]:
    """Find all .docx files in a directory."""
    if not directory.exists():
        directory.mkdir(parents=True, exist_ok=True)
        return []
    return list(directory.glob("*.docx"))

def main():
    # Define default directories
    base_dir = Path(os.path.dirname(os.path.abspath(__file__)))
    input_dir = base_dir / "input"
    target_dir = input_dir / "target"
    same_origin_dir = input_dir / "same_origin"
    reference_dir = input_dir / "reference"
    output_dir = base_dir / "output"  # Update output directory to be within project

    # Create directories if they don't exist
    for dir_path in [target_dir, same_origin_dir, reference_dir, output_dir]:
        dir_path.mkdir(parents=True, exist_ok=True)

    # Find .docx files in each directory
    target_files = find_docx_files(target_dir)
    same_origin_files = find_docx_files(same_origin_dir)
    reference_files = find_docx_files(reference_dir)

    if not target_files:
        print(f"\nNo target files found in {target_dir}")
        print("Please place your target Word document in the 'input/target' directory.")
        sys.exit(1)

    if not same_origin_files:
        print(f"\nNo same-origin files found in {same_origin_dir}")
        print("Please place your suspected same-origin Word document in the 'input/same_origin' directory.")
        sys.exit(1)

    if not reference_files:
        print(f"\nNo reference files found in {reference_dir}")
        print("Please place your reference Word document in the 'input/reference' directory.")
        sys.exit(1)

    # Use the first file from each directory
    target_file = target_files[0]
    same_origin_file = same_origin_files[0]
    reference_file = reference_files[0]

    print("\nAnalyzing the following files:")
    print(f"Target file: {target_file.name}")
    print(f"Same-origin file: {same_origin_file.name}")
    print(f"Reference file: {reference_file.name}")

    try:
        analyzer = WordForensicAnalyzer(
            str(target_file),
            str(same_origin_file),
            str(reference_file)
        )
        analyzer.validate_files()
        report, summary = analyzer.generate_report()
        
        # Save JSON report
        output_json = analyzer.output_dir / f"{Path(target_file).stem}_analysis.json"
        with open(output_json, 'w') as f:
            json.dump(report, f, indent=2)
        
        # Save summary
        output_txt = analyzer.output_dir / f"{Path(target_file).stem}_summary.txt"
        with open(output_txt, 'w') as f:
            f.write(summary)
        
        print("\nAnalysis complete!")
        print("Results saved to:")
        print(f"  - JSON report: {output_json}")
        print(f"  - Summary: {output_txt}")
    except Exception as e:
        print(f"\nError: {str(e)}")
        sys.exit(1)

if __name__ == "__main__":
    main()
