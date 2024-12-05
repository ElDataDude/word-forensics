"""
Comprehensive Document Forensics Analyzer.

This module provides tools for analyzing Word documents across multiple dimensions:
1. Metadata Analysis
2. Content Structure Analysis 
3. OOXML Analysis
4. Binary Analysis

Each dimension provides statistical patterns that can help identify document origins.
"""

import os
from pathlib import Path
import logging
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import olefile
import xml.etree.ElementTree as ET
import hashlib
import itertools
import numpy as np
from collections import Counter, defaultdict
from typing import Dict, List, Tuple, Set, Any
import json

logger = logging.getLogger(__name__)

class DocumentForensicsAnalyzer:
    """Analyzes Word documents across multiple forensic dimensions."""
    
    def __init__(self):
        """Initialize the analyzer with empty reference statistics."""
        self.reference_stats = {}
        self.field_frequencies = defaultdict(float)
        self.total_pairs = 0
        
    def analyze_reference_set(self, reference_files: List[str]) -> Dict:
        """
        Analyze a set of reference documents to establish baseline patterns.
        
        Args:
            reference_files: List of paths to reference documents
            
        Returns:
            Dictionary containing reference set statistics
        """
        all_characteristics = []
        
        # Extract characteristics from each reference file
        for file_path in reference_files:
            try:
                characteristics = self._extract_all_characteristics(file_path)
                if characteristics:  # Only add if we got valid characteristics
                    all_characteristics.append(characteristics)
            except Exception as e:
                logger.warning(f"Error analyzing {file_path}: {e}")
                
        if not all_characteristics:
            raise ValueError("No valid reference files could be analyzed")
                
        # Calculate pairwise statistics
        self.total_pairs = 0
        field_matches = defaultdict(int)
        
        # Get all possible fields from the characteristics
        all_fields = set()
        for chars in all_characteristics:
            all_fields.update(chars.keys())
        
        for char1, char2 in itertools.combinations(all_characteristics, 2):
            self.total_pairs += 1
            for field in all_fields:
                if self._field_matches(char1.get(field), char2.get(field)):
                    field_matches[field] += 1
                    
        # Calculate field frequencies
        self.field_frequencies = {
            field: matches / self.total_pairs 
            for field, matches in field_matches.items()
        }
        
        # Identify notable patterns
        always_matching = [
            field for field, freq in self.field_frequencies.items() 
            if freq > 0.99
        ]
        never_matching = [
            field for field, freq in self.field_frequencies.items() 
            if freq < 0.01 and field in field_matches  # Only count fields we've seen
        ]
        
        self.reference_stats = {
            'total_pairs': self.total_pairs,
            'field_frequencies': self.field_frequencies,
            'always_matching': always_matching,
            'never_matching': never_matching
        }
        
        return self.reference_stats

    def analyze_target_pair(self, file1: str, file2: str) -> Dict:
        """
        Analyze two documents to determine likelihood of shared origin.
        
        Args:
            file1: Path to first document
            file2: Path to second document
            
        Returns:
            Analysis results including confidence score and supporting evidence
        """
        if not self.reference_stats:
            raise ValueError("Must analyze reference set before comparing pairs")
            
        # Extract characteristics
        try:
            char1 = self._extract_all_characteristics(file1)
            char2 = self._extract_all_characteristics(file2)
        except Exception as e:
            raise ValueError(f"Error extracting characteristics: {e}")
        
        # Find matching characteristics
        matching_fields = []
        suspicious_matches = []
        anomalies = []
        
        # Get all possible fields
        all_fields = set(char1.keys()) | set(char2.keys())
        
        for field in all_fields:
            if self._field_matches(char1.get(field), char2.get(field)):
                matching_fields.append(field)
                
                # Check if this match is suspicious (rare in reference set)
                freq = self.field_frequencies.get(field, 0)
                if freq < 0.05 and char1.get(field):  # Only flag non-empty values
                    suspicious_matches.append({
                        'field': field,
                        'value': str(char1[field]),  # Convert to string for safety
                        'frequency': freq
                    })
                    
        # Calculate confidence score
        confidence = self._calculate_confidence_score(
            matching_fields, suspicious_matches, char1, char2)
            
        # Generate summary
        summary = self._generate_summary(
            confidence, matching_fields, suspicious_matches, anomalies)
            
        return {
            'confidence_score': confidence,
            'analysis_summary': summary,
            'matching_fields': matching_fields,
            'suspicious_matches': suspicious_matches,
            'anomalies': anomalies
        }
        
    def _extract_all_characteristics(self, file_path: str) -> Dict:
        """Extract all forensic characteristics from a document."""
        characteristics = {}
        
        try:
            # 1. Metadata Analysis
            metadata = self._extract_metadata(file_path)
            if metadata:
                characteristics.update(metadata)
            
            # 2. Content Structure Analysis
            structure = self._analyze_content_structure(file_path)
            if structure:
                characteristics.update(structure)
            
            # 3. OOXML Analysis
            ooxml = self._analyze_ooxml(file_path)
            if ooxml:
                characteristics.update(ooxml)
            
            # 4. Binary Analysis
            binary = self._analyze_binary(file_path)
            if binary:
                characteristics.update(binary)
                
        except Exception as e:
            logger.warning(f"Error extracting characteristics from {file_path}: {e}")
            
        return characteristics
        
    def _extract_metadata(self, file_path: str) -> Dict:
        """Extract document metadata characteristics."""
        try:
            doc = Document(file_path)
            props = doc.core_properties
            
            metadata = {
                'creator': str(getattr(props, 'author', '')),
                'last_modified_by': str(getattr(props, 'last_modified_by', '')),
                'company': str(getattr(props, 'company', '')),
                'revision_number': str(getattr(props, 'revision', '')),
                'template': '',  # Handle template separately
                'application': 'Microsoft Office Word',
                'application_version': str(getattr(props, 'version', '16.0')),
                'total_editing_time': str(getattr(props, 'total_editing_time', '0'))
            }
            
            # Try to get template info safely
            try:
                if hasattr(doc, 'settings') and doc.settings:
                    metadata['template'] = str(getattr(doc.settings, 'template', ''))
            except:
                metadata['template'] = ''
            
            # Add custom properties if any
            custom_props = self._extract_custom_properties(doc)
            if custom_props:
                metadata['custom_properties'] = custom_props
                
            return {'metadata_' + k: v for k, v in metadata.items()}
            
        except Exception as e:
            logger.warning(f"Error extracting metadata: {e}")
            return {}
        
    def _analyze_content_structure(self, file_path: str) -> Dict:
        """Analyze document content structure characteristics."""
        try:
            doc = Document(file_path)
            
            # Analyze styles
            style_counts = Counter()
            for paragraph in doc.paragraphs:
                if hasattr(paragraph, 'style') and paragraph.style:
                    style_name = getattr(paragraph.style, 'name', 'default')
                    style_counts[style_name] += 1
                    
            # Analyze sections
            section_props = []
            for section in doc.sections:
                props = {
                    'page_height': getattr(section, 'page_height', 0),
                    'page_width': getattr(section, 'page_width', 0),
                    'left_margin': getattr(section, 'left_margin', 0),
                    'right_margin': getattr(section, 'right_margin', 0),
                    'top_margin': getattr(section, 'top_margin', 0),
                    'bottom_margin': getattr(section, 'bottom_margin', 0),
                    'header_distance': getattr(section, 'header_distance', 0),
                    'footer_distance': getattr(section, 'footer_distance', 0),
                    'orientation': getattr(section, 'orientation', 0)
                }
                section_props.append(props)
                
            # Analyze tables
            table_structures = []
            for table in doc.tables:
                structure = {
                    'rows': len(table.rows),
                    'cols': len(table.columns)
                }
                if hasattr(table, 'style') and table.style:
                    structure['style'] = getattr(table.style, 'name', 'default')
                table_structures.append(structure)
                
            return {
                'structure_styles': dict(style_counts),
                'structure_sections': section_props,
                'structure_tables': table_structures
            }
            
        except Exception as e:
            logger.warning(f"Error analyzing content structure: {e}")
            return {}
        
    def _analyze_ooxml(self, file_path: str) -> Dict:
        """Analyze OOXML characteristics."""
        try:
            doc = Document(file_path)
            
            # Analyze XML namespaces
            namespaces = set()
            try:
                for elem in doc._element.iter():
                    if hasattr(elem, 'tag') and isinstance(elem.tag, str):
                        if elem.tag.startswith('{'):
                            ns = elem.tag[1:].split('}')[0]
                            namespaces.add(ns)
            except:
                pass
                    
            # Analyze relationships
            rels = set()
            try:
                if hasattr(doc.part, 'rels'):
                    for rel in doc.part.rels.values():
                        if hasattr(rel, 'reltype'):
                            rels.add(rel.reltype)
            except:
                pass
                
            # Analyze content types
            content_types = set()
            try:
                if hasattr(doc.part, 'blob'):
                    tree = ET.parse(doc.part.blob)
                    for elem in tree.iter():
                        if 'ContentType' in elem.attrib:
                            content_types.add(elem.attrib['ContentType'])
            except:
                pass
                
            return {
                'ooxml_namespaces': sorted(list(namespaces)),
                'ooxml_relationships': sorted(list(rels)),
                'ooxml_content_types': sorted(list(content_types))
            }
            
        except Exception as e:
            logger.warning(f"Error analyzing OOXML: {e}")
            return {}
        
    def _analyze_binary(self, file_path: str) -> Dict:
        """Analyze binary file characteristics."""
        try:
            # Read file signature
            with open(file_path, 'rb') as f:
                header = f.read(8)
                file_sig = header.hex()
                
            # Analyze OLE structure if possible
            ole_streams = []
            if olefile.isOleFile(file_path):
                ole = olefile.OleFileIO(file_path)
                ole_streams = sorted(ole.listdir())
                ole.close()
                
            # Calculate various hashes
            hash_md5 = hashlib.md5()
            hash_sha1 = hashlib.sha1()
            
            with open(file_path, 'rb') as f:
                while chunk := f.read(8192):
                    hash_md5.update(chunk)
                    hash_sha1.update(chunk)
                    
            return {
                'binary_signature': file_sig,
                'binary_ole_streams': ole_streams,
                'binary_md5': hash_md5.hexdigest(),
                'binary_sha1': hash_sha1.hexdigest()
            }
            
        except Exception as e:
            logger.warning(f"Error analyzing binary: {e}")
            return {}
        
    def _extract_custom_properties(self, doc: Document) -> Dict:
        """Extract custom document properties."""
        custom_props = {}
        try:
            if hasattr(doc, 'custom_properties'):
                for prop in doc.custom_properties:
                    if hasattr(prop, 'name') and hasattr(prop, 'value'):
                        custom_props[prop.name] = str(prop.value)
        except:
            pass
        return custom_props
        
    def _get_all_fields(self) -> List[str]:
        """Get list of all characteristic fields being analyzed."""
        return [
            # Metadata fields
            'metadata_creator', 'metadata_last_modified_by', 'metadata_company',
            'metadata_revision_number', 'metadata_template', 'metadata_application',
            'metadata_application_version', 'metadata_total_editing_time',
            'metadata_custom_properties',
            
            # Structure fields
            'structure_styles', 'structure_sections', 'structure_tables',
            
            # OOXML fields
            'ooxml_namespaces', 'ooxml_relationships', 'ooxml_content_types',
            
            # Binary fields
            'binary_signature', 'binary_ole_streams', 'binary_md5', 'binary_sha1'
        ]
        
    def _field_matches(self, value1: Any, value2: Any) -> bool:
        """Check if two field values match, handling different types appropriately."""
        if value1 is None or value2 is None:
            return False
            
        # Convert to JSON strings for comparison of complex types
        if isinstance(value1, (dict, list, set)):
            try:
                return json.dumps(value1, sort_keys=True) == json.dumps(value2, sort_keys=True)
            except:
                return False
        else:
            # For simple types, direct comparison
            return value1 == value2
            
    def _calculate_confidence_score(
        self, matching_fields: List[str], 
        suspicious_matches: List[Dict],
        char1: Dict, char2: Dict
    ) -> float:
        """Calculate confidence score for shared origin."""
        if not matching_fields:
            return 0.0
            
        # Weight different types of matches
        weights = {
            'metadata_': 0.3,    # Metadata matches
            'structure_': 0.3,   # Content structure matches
            'ooxml_': 0.2,      # OOXML characteristic matches
            'binary_': 0.2      # Binary characteristic matches
        }
        
        # Calculate weighted score
        total_weight = 0
        weighted_score = 0
        
        for field in matching_fields:
            # Determine which category this field belongs to
            category = next(
                (cat for cat in weights.keys() if field.startswith(cat)), 
                'other_'
            )
            
            # Get the weight for this category
            weight = weights.get(category, 0.1)
            
            # Add to weighted score
            total_weight += weight
            weighted_score += weight
            
        # Normalize score
        base_score = weighted_score / total_weight if total_weight > 0 else 0
        
        # Adjust for suspicious matches
        suspicion_bonus = len(suspicious_matches) * 0.1
        
        return min(1.0, base_score + suspicion_bonus)
        
    def _generate_summary(
        self, confidence: float, 
        matching_fields: List[str],
        suspicious_matches: List[Dict],
        anomalies: List[Dict]
    ) -> str:
        """Generate human-readable analysis summary."""
        # Determine confidence level
        if confidence > 0.8:
            confidence_text = "HIGH"
        elif confidence > 0.5:
            confidence_text = "MEDIUM"
        else:
            confidence_text = "LOW"
            
        summary = [
            f"Overall Assessment: {confidence_text} probability of shared origin",
            f"Confidence Score: {confidence:.2f}\n"
        ]
        
        if matching_fields:
            summary.append("Matching Fields:")
            for field in matching_fields:
                summary.append(f"- {field}")
            summary.append("")
            
        if suspicious_matches:
            summary.append("Suspicious Matches:")
            for match in suspicious_matches:
                summary.append(
                    f"- {match['field']}: {match['value']} "
                    f"(occurs in {match['frequency']:.1%} of reference pairs)"
                )
                
        if anomalies:
            summary.append("\nAnomalies Detected:")
            for anomaly in anomalies:
                summary.append(f"- {anomaly['description']}")
                
        return "\n".join(summary)
